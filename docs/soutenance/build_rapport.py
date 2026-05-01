"""
build_rapport.py — Génère le rapport de soutenance en .docx.

Structure type d'un mémoire de Licence : couverture, remerciements, sommaire,
introduction, méthodologie, 5 chapitres (un par phase), conclusion, références,
annexes.

Dépendance : python-docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).parent
OUTPUT = ROOT / "rapport_soutenance.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
BLUE = RGBColor(0x0D, 0x6E, 0xFD)
GREY = RGBColor(0x6C, 0x75, 0x7D)
DARK = RGBColor(0x1A, 0x23, 0x32)


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str = "CCCCCC") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r2 = paragraph.add_run()
    r2.font.name = "Arial"; r2.font.size = Pt(9)
    r2._r.append(f1); r2._r.append(instr); r2._r.append(f2)


def add_toc(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Mettez à jour la table des matières dans Word (clic droit, Mettre à jour)."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, instr, f2, txt, f3):
        run._r.append(el)


def styled(doc, text, *, style=None, align=None, bold=False, italic=False,
           color=None, size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Arial"
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.color.rgb = BLUE
            run.font.size = Pt(20)
        elif level == 2:
            run.font.color.rgb = DARK
            run.font.size = Pt(14)
        else:
            run.font.color.rgb = DARK
            run.font.size = Pt(12)
    return h


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.name = "Arial"
        r.font.size = Pt(11)


def numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(it)
        r.font.name = "Arial"
        r.font.size = Pt(11)


def table(doc, headers, rows, widths_cm):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.width = Cm(widths_cm[j])
        set_cell_shading(c, "D5E8F0")
        set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ""
        rr = c.paragraphs[0].add_run(h)
        rr.font.name = "Arial"; rr.font.size = Pt(10); rr.bold = True
    for i, row in enumerate(rows, start=1):
        for j, txt in enumerate(row):
            c = t.cell(i, j)
            c.width = Cm(widths_cm[j])
            set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.paragraphs[0].text = ""
            rr = c.paragraphs[0].add_run(txt)
            rr.font.name = "Arial"; rr.font.size = Pt(10)
    doc.add_paragraph()


def code_block(doc, code: str, lang: str = ""):
    """Bloc de code monospace dans une cellule grisée."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    c = t.cell(0, 0)
    c.width = Cm(16)
    set_cell_shading(c, "F5F5F5")
    set_cell_borders(c, "DDDDDD")
    c.paragraphs[0].text = ""
    for line in code.split("\n"):
        p = c.add_paragraph() if line != code.split("\n")[0] else c.paragraphs[0]
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
def build():
    doc = Document()

    # Styles globaux
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    sec.different_first_page_header_footer = True

    # ---------- COUVERTURE ----------
    for _ in range(3):
        doc.add_paragraph()
    styled(doc, "Université Iba Der Thiam de Thiès",
           align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    styled(doc, "UFR Sciences de l'Ingénieur",
           align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    styled(doc, "Département de Géomatique",
           align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
    for _ in range(4):
        doc.add_paragraph()
    styled(doc, "RAPPORT DE PROJET",
           align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
           color=GREY)
    styled(doc, "Programmation Orientée Objet",
           align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12,
           color=GREY)
    for _ in range(2):
        doc.add_paragraph()
    styled(doc, "Gestion de l'Emprunt du Matériel",
           align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24,
           color=BLUE)
    styled(doc, "Topographique et Géodésique",
           align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20,
           color=BLUE)
    styled(doc, "avec assistant IA fine-tuné",
           align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=14)
    for _ in range(5):
        doc.add_paragraph()
    styled(doc, "Présenté par",
           align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
    styled(doc, "Bineta Elimane HANNE  &  Aminata KOUNTA",
           align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    styled(doc, "Étudiantes en Licence 2 Géomatique",
           align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
    for _ in range(4):
        doc.add_paragraph()
    styled(doc, "Année universitaire 2025 — 2026",
           align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)

    # ---------- REMERCIEMENTS ----------
    doc.add_page_break()
    heading(doc, "Remerciements", level=1)
    para(doc,
        "Au terme de ce projet, nous tenons à exprimer notre reconnaissance "
        "à toutes les personnes qui ont contribué à sa réalisation."
    )
    para(doc,
        "Nous remercions l'équipe pédagogique de l'UFR Sciences de l'Ingénieur "
        "de l'Université Iba Der Thiam de Thiès pour la qualité de leur "
        "enseignement en programmation orientée objet, en bases de données "
        "et en topographie, sans lequel ce projet n'aurait pas vu le jour."
    )
    para(doc,
        "Nous remercions également les techniciens du laboratoire de topographie "
        "qui nous ont aidées à comprendre le fonctionnement réel du parc matériel "
        "de l'UFR, ainsi qu'à identifier les besoins concrets auxquels notre "
        "application devait répondre."
    )
    para(doc,
        "Enfin, nous remercions nos camarades de promotion pour leurs retours "
        "lors des phases de test, et nos familles pour leur soutien constant."
    )

    # ---------- RÉSUMÉ ----------
    doc.add_page_break()
    heading(doc, "Résumé", level=1)
    para(doc,
        "Ce projet présente la conception et le développement d'une application "
        "web complète pour la gestion de l'emprunt du matériel topographique "
        "et géodésique de l'UFR Sciences de l'Ingénieur. Réalisée en binôme "
        "dans le cadre du cours de Programmation Orientée Objet de Licence 2 "
        "Géomatique, elle remplace la gestion papier actuelle par un système "
        "traçable, accessible 24 h / 24, et enrichi d'un chatbot intelligent "
        "fine-tuné sur le matériel de l'UFR."
    )
    para(doc,
        "L'application est construite sur Django et Django REST Framework côté "
        "backend, Bootstrap 5 et Leaflet.js côté frontend, avec un microservice "
        "FastAPI dédié à l'inférence d'un modèle Phi-3-mini adapté par LoRA. "
        "Elle implémente les concepts fondamentaux de l'orientation objet : "
        "héritage, encapsulation, polymorphisme et abstraction."
    )
    styled(doc, "Mots-clés : ", bold=True)
    p = doc.paragraphs[-1]
    r = p.add_run(
        "Django, REST API, POO, fine-tuning, LoRA, chatbot, topographie, "
        "géolocalisation, Leaflet, UFR SI, Thiès."
    )
    r.font.name = "Arial"; r.font.size = Pt(11); r.italic = True

    # ---------- TABLE DES MATIÈRES ----------
    doc.add_page_break()
    heading(doc, "Table des matières", level=1)
    add_toc(doc)

    # ---------- HEADER / FOOTER ----------
    header = sec.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]; hp.text = ""
    r = hp.add_run("Gestion de l'emprunt du matériel — UFR SI")
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.color.rgb = GREY
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = sec.footer
    fp = footer.paragraphs[0]; fp.text = ""; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(fp)

    sec.first_page_header.paragraphs[0].text = ""
    sec.first_page_footer.paragraphs[0].text = ""

    # =====================================================================
    # INTRODUCTION
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Introduction générale", level=1)

    heading(doc, "Contexte", level=2)
    para(doc,
        "L'UFR Sciences de l'Ingénieur de l'Université Iba Der Thiam de Thiès "
        "met à disposition de ses étudiants un parc important de matériel "
        "topographique et géodésique : stations totales, récepteurs GNSS, "
        "niveaux optiques et électroniques, accessoires (trépieds, prismes, "
        "embases), équipements informatiques mobiles. Ce matériel, dont la "
        "valeur dépasse plusieurs dizaines de millions de francs CFA, est "
        "indispensable à la formation pratique en topographie."
    )
    para(doc,
        "Sa gestion repose actuellement sur un système papier : cahier "
        "d'emprunt, fiches manuelles de retour, suivi assuré par les "
        "techniciens. Ce mode de gestion présente plusieurs limites : "
        "traçabilité difficile, doubles attributions involontaires, retards "
        "non détectés, perte d'historique, absence d'indicateurs."
    )

    heading(doc, "Problématique", level=2)
    para(doc,
        "Comment digitaliser et fluidifier la gestion du matériel topographique "
        "de l'UFR pour la rendre transparente et accessible 24 h / 24, tout en "
        "restant simple à utiliser pour les étudiants et utile pour les "
        "techniciens et enseignants ?"
    )
    para(doc,
        "À cette problématique s'ajoute une exigence pédagogique : ce projet "
        "doit illustrer concrètement les concepts de la programmation orientée "
        "objet enseignés en cours, et démontrer notre capacité à conduire un "
        "projet logiciel complet, du cahier des charges à la mise en service."
    )

    heading(doc, "Objectifs", level=2)
    bullets(doc, [
        "Concevoir et développer une application web couvrant le cycle complet "
        "de l'emprunt : demande, validation, sortie, restitution.",
        "Intégrer une géolocalisation des sorties terrain via une carte "
        "interactive centrée sur Thiès.",
        "Développer un chatbot fine-tuné, capable de répondre aux questions "
        "techniques et procédurales des étudiants.",
        "Fournir un tableau de bord administrateur pour le suivi en temps réel.",
        "Documenter rigoureusement chaque phase, en lien avec les concepts POO.",
    ])

    heading(doc, "Plan du rapport", level=2)
    para(doc,
        "Ce rapport est structuré en cinq chapitres correspondant aux cinq "
        "phases du projet. Le chapitre 1 présente la conception (UML, MCD/MLD, "
        "architecture). Le chapitre 2 décrit l'implémentation backend avec "
        "Django et l'API REST. Le chapitre 3 détaille le frontend Bootstrap "
        "et l'intégration cartographique Leaflet. Le chapitre 4 expose la "
        "construction du corpus et du dataset d'entraînement. Le chapitre 5 "
        "présente le fine-tuning LoRA et l'intégration du chatbot. Une "
        "conclusion résume les acquis et ouvre sur les perspectives."
    )

    heading(doc, "Organisation du binôme", level=2)
    table(doc,
        headers=["Membre", "Périmètre principal", "Responsabilités techniques"],
        rows=[
            ["Bineta Elimane Hanne",
             "Backend & IA",
             "Modèles Django, API REST, services métier, dataset Q/R, "
             "fine-tuning LoRA, microservice FastAPI"],
            ["Aminata Kounta",
             "Frontend & Cartographie",
             "Templates Bootstrap, navigation, formulaires, "
             "carte Leaflet, intégration UX, manuel utilisateur"],
        ],
        widths_cm=[3.5, 3.5, 9.0],
    )
    para(doc,
        "Cette répartition a permis un travail en parallèle sur deux branches "
        "Git distinctes (feat/backend, feat/frontend, feat/ia), avec des "
        "synchronisations régulières par fusion sur la branche main."
    )

    # =====================================================================
    # CHAPITRE 1 — CONCEPTION
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Chapitre 1 — Conception", level=1)

    heading(doc, "1.1 Démarche méthodologique", level=2)
    para(doc,
        "Avant toute écriture de code, nous avons mené une phase complète de "
        "conception inspirée de la démarche Merise et de la modélisation UML. "
        "L'objectif était double : poser une vision claire du système à "
        "construire, et identifier les classes métier qui structureraient "
        "ensuite l'implémentation orientée objet."
    )

    heading(doc, "1.2 Étude des besoins", level=2)
    para(doc,
        "Quatre acteurs principaux ont été identifiés à partir de l'analyse "
        "des pratiques actuelles et d'entretiens informels avec les "
        "techniciens du laboratoire :"
    )
    table(doc,
        headers=["Acteur", "Besoins principaux"],
        rows=[
            ["Étudiant", "Consulter le matériel, soumettre des demandes, "
             "suivre leur statut, poser des questions techniques"],
            ["Enseignant", "Valider ou refuser les demandes, voir les "
             "sorties de sa filière"],
            ["Technicien", "Marquer les sorties et les retours, ouvrir des "
             "fiches de maintenance, mettre à jour le stock"],
            ["Administrateur", "Gérer les utilisateurs, le catalogue, "
             "consulter les statistiques globales"],
        ],
        widths_cm=[4.0, 12.0],
    )

    heading(doc, "1.3 Diagramme de cas d'utilisation", level=2)
    para(doc,
        "Le diagramme de cas d'utilisation a été modélisé en UML, mettant en "
        "évidence les interactions entre les quatre acteurs et le système. "
        "Les cas principaux sont : s'inscrire, se connecter, consulter le "
        "catalogue, soumettre une demande, valider une demande, restituer "
        "le matériel, gérer la maintenance, et discuter avec le chatbot."
    )
    para(doc,
        "Le détail figure dans le fichier docs/conception/01_cas_utilisation.md "
        "du dépôt, sous forme d'un diagramme Mermaid rendu directement par GitHub."
    )

    heading(doc, "1.4 Diagramme de classes", level=2)
    para(doc,
        "Le diagramme de classes décompose le système en six entités principales :"
    )
    bullets(doc, [
        "Utilisateur : hérite de AbstractUser de Django, ajoute le rôle "
        "(Étudiant, Enseignant, Technicien, Administrateur), la filière, le niveau.",
        "Categorie : regroupe le matériel par famille (Stations totales, GNSS, etc.).",
        "Materiel : décrit chaque équipement avec sa marque, son modèle, son état "
        "et ses quantités totale et disponible.",
        "Demande : représente une demande d'emprunt avec ses dates, son motif, "
        "son statut et son emplacement.",
        "LigneDemande : table de liaison entre Demande et Materiel avec une quantité.",
        "ConversationChat : conserve l'historique des échanges entre un "
        "utilisateur et le chatbot IA.",
    ])
    para(doc,
        "Les relations principales sont : Utilisateur 1—N Demande, "
        "Demande 1—N LigneDemande, Materiel 1—N LigneDemande, "
        "Categorie 1—N Materiel, Demande 1—1 Emplacement (géolocalisation), "
        "Demande 1—1 Restitution."
    )

    heading(doc, "1.5 Modèle Conceptuel et Logique de Données (MCD/MLD)", level=2)
    para(doc,
        "Conformément à la méthode Merise enseignée en cours, nous avons "
        "transformé le modèle conceptuel en modèle logique avec les "
        "contraintes suivantes :"
    )
    bullets(doc, [
        "Une demande ne peut pas avoir une date_fin antérieure à sa date_debut "
        "(contrainte de cohérence temporelle, implémentée par CheckConstraint).",
        "Une ligne de demande est unique pour le couple (demande, matériel).",
        "Un emplacement est lié à exactement une demande (relation OneToOne).",
        "Une restitution est unique pour une demande donnée.",
    ])

    heading(doc, "1.6 Architecture technique", level=2)
    para(doc,
        "L'architecture retenue est en trois couches, avec une séparation "
        "stricte des responsabilités :"
    )
    table(doc,
        headers=["Couche", "Composant", "Technologie"],
        rows=[
            ["Présentation", "Pages HTML, formulaires, carte interactive",
             "Bootstrap 5 + Leaflet.js"],
            ["Application", "Logique métier, API REST, authentification",
             "Django + DRF + JWT"],
            ["Données", "Persistance",
             "SQLite (dev) / PostgreSQL+PostGIS (prod)"],
            ["IA", "Inférence du modèle",
             "FastAPI + Phi-3-mini + LoRA"],
        ],
        widths_cm=[3.5, 6.5, 6.0],
    )

    heading(doc, "1.7 Choix technologiques justifiés", level=2)
    para(doc,
        "Django a été choisi pour sa robustesse, sa documentation et son "
        "écosystème (admin auto-généré, ORM puissant, sécurité par défaut). "
        "DRF complète Django avec une API REST mature et des fonctionnalités "
        "(serializers, viewsets, permissions) qui alignent bien avec la POO."
    )
    para(doc,
        "Bootstrap 5 a été retenu pour le frontend car il offre un design "
        "responsive directement utilisable, avec des composants riches et "
        "une grande communauté. Leaflet.js, plus léger que Mapbox ou Google "
        "Maps, est libre et s'intègre nativement avec OpenStreetMap, ce qui "
        "convient parfaitement à la cartographie du Sénégal."
    )
    para(doc,
        "Pour l'IA, le choix s'est porté sur Phi-3-mini (3,8 milliards de "
        "paramètres) plutôt que des modèles plus grands (Llama 7B, Mistral 7B) "
        "afin de pouvoir tourner sur du matériel modeste. La méthode LoRA "
        "permet d'adapter ce modèle avec très peu de ressources GPU."
    )

    # =====================================================================
    # CHAPITRE 2 — BACKEND
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Chapitre 2 — Backend Django", level=1)

    heading(doc, "2.1 Organisation du projet Django", level=2)
    para(doc,
        "Le backend est découpé en cinq applications Django, chacune ayant "
        "une responsabilité unique :"
    )
    table(doc,
        headers=["Application", "Responsabilité"],
        rows=[
            ["users", "Modèle Utilisateur (héritage de AbstractUser), gestion des rôles"],
            ["materiel", "Catégories, matériels, fiches de maintenance"],
            ["emprunts", "Demandes, lignes de demande, emplacements, restitutions"],
            ["chatbot", "Conversations, service IA"],
            ["api", "Routeur DRF central regroupant tous les endpoints REST"],
            ["pages", "Vues HTML pour les pages côté étudiant"],
        ],
        widths_cm=[3.5, 12.5],
    )

    heading(doc, "2.2 Modèle Utilisateur — héritage", level=2)
    para(doc,
        "Le modèle Utilisateur illustre l'un des concepts fondamentaux de la POO : "
        "l'héritage. Plutôt que de réinventer un système d'authentification, nous "
        "héritons de AbstractUser fourni par Django, qui propose déjà username, "
        "password, email, first_name, last_name, et la gestion des permissions."
    )
    code_block(doc,
        "class Role(models.TextChoices):\n"
        "    ETUDIANT = 'ETUDIANT', 'Étudiant'\n"
        "    ENSEIGNANT = 'ENSEIGNANT', 'Enseignant'\n"
        "    TECHNICIEN = 'TECHNICIEN', 'Technicien'\n"
        "    ADMINISTRATEUR = 'ADMINISTRATEUR', 'Administrateur'\n\n"
        "class Utilisateur(AbstractUser):\n"
        "    role = models.CharField(max_length=20, choices=Role.choices,\n"
        "                            default=Role.ETUDIANT)\n"
        "    filiere = models.CharField(max_length=100, blank=True)\n"
        "    niveau = models.CharField(max_length=20, blank=True)\n"
        "    telephone = models.CharField(max_length=20, blank=True)\n\n"
        "    def est_administrateur(self) -> bool:\n"
        "        return self.role == Role.ADMINISTRATEUR\n\n"
        "    def peut_emprunter(self) -> bool:\n"
        "        return self.is_active and self.role != Role.ADMINISTRATEUR"
    )
    para(doc,
        "Cette approche illustre aussi l'encapsulation : la logique métier "
        "(« qui peut emprunter ? ») est encapsulée dans des méthodes du modèle "
        "plutôt que dispersée dans les vues."
    )

    heading(doc, "2.3 Modèle Materiel — encapsulation et états", level=2)
    para(doc,
        "Le modèle Materiel encapsule la gestion d'état avec une énumération "
        "(EtatMateriel : DISPONIBLE, EMPRUNTE, EN_MAINTENANCE, HORS_SERVICE). "
        "Les transitions d'état passent par des méthodes dédiées qui appliquent "
        "les invariants métier."
    )
    code_block(doc,
        "class Materiel(models.Model):\n"
        "    nom = models.CharField(max_length=200)\n"
        "    categorie = models.ForeignKey(Categorie, on_delete=models.PROTECT)\n"
        "    etat = models.CharField(max_length=20, choices=EtatMateriel.choices,\n"
        "                            default=EtatMateriel.DISPONIBLE)\n"
        "    quantite_totale = models.PositiveIntegerField(default=1)\n"
        "    quantite_disponible = models.PositiveIntegerField(default=1)\n\n"
        "    def est_disponible(self) -> bool:\n"
        "        return (self.etat == EtatMateriel.DISPONIBLE\n"
        "                and self.quantite_disponible > 0)\n\n"
        "    def marquer_en_panne(self):\n"
        "        self.etat = EtatMateriel.EN_MAINTENANCE\n"
        "        self.save()\n\n"
        "    def remettre_en_service(self):\n"
        "        self.etat = EtatMateriel.DISPONIBLE\n"
        "        self.save()"
    )

    heading(doc, "2.4 Modèle Demande — polymorphisme par méthodes d'action", level=2)
    para(doc,
        "Le modèle Demande implémente une machine à états : une demande passe "
        "successivement par les statuts EN_ATTENTE, APPROUVEE, EN_COURS, "
        "RESTITUEE (ou REFUSEE / ANNULEE). Chaque transition est portée par "
        "une méthode dédiée, ce qui rend le code lisible et facilement testable."
    )
    code_block(doc,
        "class Demande(models.Model):\n"
        "    utilisateur = models.ForeignKey(Utilisateur, related_name='demandes')\n"
        "    date_debut = models.DateField()\n"
        "    date_fin = models.DateField()\n"
        "    motif = models.TextField(blank=True)\n"
        "    statut = models.CharField(choices=StatutDemande.choices,\n"
        "                              default=StatutDemande.EN_ATTENTE)\n\n"
        "    class Meta:\n"
        "        constraints = [\n"
        "            CheckConstraint(check=Q(date_fin__gte=F('date_debut')),\n"
        "                            name='date_fin_apres_debut'),\n"
        "        ]\n\n"
        "    def valider(self, valideur): ...\n"
        "    def refuser(self, motif_refus): ...\n"
        "    def marquer_en_cours(self): ...\n"
        "    def annuler(self): ..."
    )

    heading(doc, "2.5 API REST avec Django REST Framework", level=2)
    para(doc,
        "Tous les modèles métier sont exposés via une API REST construite avec "
        "DRF. Les ViewSets génèrent automatiquement les endpoints CRUD (GET, "
        "POST, PUT, DELETE), tandis que les Serializers gèrent la conversion "
        "objet Python ↔ JSON. Un routeur DRF central regroupe tous les "
        "endpoints à /api/."
    )
    table(doc,
        headers=["Endpoint", "Méthode", "Description"],
        rows=[
            ["/api/utilisateurs/", "GET, POST", "Lister / créer des utilisateurs"],
            ["/api/categories/", "GET", "Lister les catégories"],
            ["/api/materiels/", "GET, POST", "Catalogue + création"],
            ["/api/demandes/", "GET, POST", "Demandes de l'utilisateur courant"],
            ["/api/demandes/{id}/valider/", "POST", "Action métier : validation"],
            ["/api/restitutions/", "POST", "Enregistrer une restitution"],
            ["/api/chat/", "GET, POST", "Conversations avec le chatbot"],
            ["/api/chat/{id}/envoyer/", "POST", "Envoyer un message"],
            ["/api/auth/login/", "POST", "Authentification JWT"],
        ],
        widths_cm=[6.0, 2.5, 7.5],
    )

    heading(doc, "2.6 Authentification et permissions", level=2)
    para(doc,
        "L'authentification utilise JWT (JSON Web Tokens) via la bibliothèque "
        "djangorestframework-simplejwt. Chaque appel API doit présenter un "
        "token signé valide. Les permissions sont gérées par classes DRF, "
        "ce qui permet de restreindre l'accès finement (par exemple : "
        "seul un Enseignant peut valider une demande)."
    )

    heading(doc, "2.7 Données initiales (fixtures)", level=2)
    para(doc,
        "À partir de l'inventaire 2023 du laboratoire UFR SI, nous avons "
        "constitué un jeu de données initial réaliste : 7 catégories et "
        "23 matériels distincts, avec leurs quantités et états réels. Ces "
        "données sont chargées via les fixtures JSON de Django."
    )

    # =====================================================================
    # CHAPITRE 3 — FRONTEND
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Chapitre 3 — Frontend & Cartographie", level=1)

    heading(doc, "3.1 Choix de l'approche", level=2)
    para(doc,
        "Plutôt qu'un SPA (single page application) avec React ou Vue, qui "
        "aurait alourdi le projet, nous avons opté pour un rendu côté serveur "
        "via les templates Django, agrémentés de JavaScript vanilla pour les "
        "interactions dynamiques. Bootstrap 5 fournit la grille responsive et "
        "les composants UI (navbar, cartes, formulaires, modals)."
    )

    heading(doc, "3.2 Structure des templates", level=2)
    para(doc,
        "Tous les templates héritent d'un fichier base.html qui définit la "
        "navbar, le footer et les blocs content / extra_head / extra_scripts. "
        "Cette approche illustre la notion d'héritage à l'échelle des templates."
    )
    table(doc,
        headers=["Template", "Rôle"],
        rows=[
            ["base.html", "Layout commun, navbar conditionnelle, footer"],
            ["home.html", "Page d'accueil avec hero et statistiques"],
            ["registration/login.html", "Connexion"],
            ["registration/register.html", "Inscription étudiant"],
            ["materiel/catalogue.html", "Grille filtrée du matériel"],
            ["emprunts/mes_emprunts.html", "Tableau de mes demandes"],
            ["emprunts/nouvelle_demande.html", "Formulaire avec carte Leaflet"],
            ["chatbot/chat.html", "Interface de discussion avec l'IA"],
        ],
        widths_cm=[6.0, 10.0],
    )

    heading(doc, "3.3 Intégration de la carte Leaflet", level=2)
    para(doc,
        "Le formulaire de nouvelle demande embarque une carte Leaflet centrée "
        "sur Thiès (latitude 14,789 ; longitude -16,926). L'utilisateur clique "
        "sur la carte pour placer un marqueur, puis peut le déplacer par "
        "glisser-déposer. Les coordonnées sont automatiquement renseignées "
        "dans les champs cachés du formulaire."
    )
    code_block(doc,
        "const THIES = [14.789, -16.926];\n"
        "const carte = L.map('carte').setView(THIES, 13);\n\n"
        "L.tileLayer('https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png', {\n"
        "  attribution: '© OpenStreetMap contributors',\n"
        "}).addTo(carte);\n\n"
        "let marqueur = null;\n"
        "carte.on('click', (e) => {\n"
        "  const { lat, lng } = e.latlng;\n"
        "  if (marqueur) marqueur.setLatLng(e.latlng);\n"
        "  else marqueur = L.marker(e.latlng, { draggable: true }).addTo(carte);\n"
        "  document.getElementById('latitude').value = lat.toFixed(6);\n"
        "  document.getElementById('longitude').value = lng.toFixed(6);\n"
        "});"
    )

    heading(doc, "3.4 Validations côté client", level=2)
    para(doc,
        "Le formulaire applique plusieurs validations en JavaScript pour "
        "améliorer l'expérience utilisateur sans attendre l'aller-retour serveur :"
    )
    bullets(doc, [
        "Date de fin >= date de début (mise à jour dynamique du min de date_fin).",
        "Quantité limitée au stock disponible affiché.",
        "Au moins un matériel coché avant de pouvoir soumettre.",
        "Compteur en temps réel des articles sélectionnés.",
    ])

    heading(doc, "3.5 Charte graphique", level=2)
    para(doc,
        "Une charte simple a été définie dans frontend/static/css/main.css : "
        "couleur primaire bleu Bootstrap (#0d6efd), accent jaune (#ffc107), "
        "fond clair (#f5f7fa), typographie Segoe UI / Roboto. Le hero de la "
        "page d'accueil utilise un dégradé linéaire bleu pour donner du "
        "caractère sans alourdir."
    )

    # =====================================================================
    # CHAPITRE 4 — CORPUS & DATASET IA
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Chapitre 4 — Corpus IA & Dataset", level=1)

    heading(doc, "4.1 RAG ou fine-tuning ?", level=2)
    para(doc,
        "Deux approches étaient envisageables pour spécialiser le chatbot :"
    )
    table(doc,
        headers=["Approche", "Avantages", "Inconvénients"],
        rows=[
            ["RAG (Retrieval-Augmented Generation)",
             "Pas d'entraînement, base de connaissances modifiable",
             "Dépendance à un LLM externe (coût API, latence, hors-ligne impossible)"],
            ["Fine-tuning",
             "Modèle autonome, hors-ligne, intégré au stack",
             "Entraînement coûteux, données figées"],
        ],
        widths_cm=[4.5, 5.5, 6.0],
    )
    para(doc,
        "Pour un projet pédagogique avec un sujet bien borné (matériel UFR + "
        "procédures fixes), nous avons opté pour le fine-tuning. Cette approche "
        "est plus instructive et démontre concrètement les concepts de "
        "machine learning enseignés en cours."
    )

    heading(doc, "4.2 Constitution du corpus", level=2)
    para(doc,
        "Le corpus de connaissances est rédigé à la main, structuré en "
        "quatre fichiers Markdown thématiques :"
    )
    table(doc,
        headers=["Fichier", "Contenu"],
        rows=[
            ["01_materiels.md", "Fiches techniques détaillées (précision, portée, usage) "
             "pour les 23 matériels du parc"],
            ["02_procedures.md", "Règles d'emprunt, statuts, restitution, cas particuliers"],
            ["03_topographie.md", "Notions techniques (angles, distances, nivellement, "
             "GNSS, projections UTM Sénégal)"],
            ["04_faq.md", "Questions fréquemment posées par les étudiants"],
        ],
        widths_cm=[4.5, 11.5],
    )

    heading(doc, "4.3 Format Alpaca", level=2)
    para(doc,
        "Le format Alpaca s'est imposé comme un standard pour le fine-tuning "
        "supervisé d'instruction-following models. Chaque exemple est un objet "
        "JSON avec trois champs : instruction (la question), input (contexte "
        "additionnel, optionnel), output (la réponse attendue)."
    )

    heading(doc, "4.4 Augmentation par paraphrase", level=2)
    para(doc,
        "À partir de 25 paires Q/R rédigées à la main et équilibrées sur les "
        "quatre catégories, un script Python applique une augmentation par "
        "paraphrase d'instruction selon le type de question (« Qu'est-ce que », "
        "« Comment », « Pourquoi »…). Le facteur d'augmentation est d'environ "
        "2,6x, portant le dataset à 65 exemples."
    )

    heading(doc, "4.5 Split stratifié", level=2)
    para(doc,
        "Plutôt qu'un split aléatoire global qui pourrait déséquilibrer les "
        "catégories, nous avons appliqué un split stratifié par catégorie, "
        "garantissant qu'au moins un exemple par catégorie figure dans le "
        "set de validation."
    )
    table(doc,
        headers=["Catégorie", "Train", "Validation", "Total"],
        rows=[
            ["materiel", "15", "4", "19"],
            ["procedures", "10", "2", "12"],
            ["topographie", "21", "5", "26"],
            ["faq", "8", "0 (pas assez)", "8"],
            ["TOTAL", "54", "11", "65"],
        ],
        widths_cm=[5.0, 3.5, 4.0, 3.5],
    )

    heading(doc, "4.6 Limites identifiées", level=2)
    bullets(doc, [
        "Taille modeste : 65 exemples sont peu pour un fine-tuning robuste. "
        "Une mise en production demanderait au moins 500 à 1000 exemples.",
        "Augmentation purement syntaxique : nous ne générons pas de paraphrases "
        "sémantiques, ce que ferait une LLM-as-augmenter.",
        "Absence de hard negatives : le dataset ne contient pas de questions "
        "hors-périmètre avec réponses du type « je ne sais pas ».",
        "Pas d'évaluation automatique : seule la perplexité serait disponible.",
    ])

    # =====================================================================
    # CHAPITRE 5 — FINE-TUNING & INTÉGRATION
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Chapitre 5 — Fine-tuning & Intégration", level=1)

    heading(doc, "5.1 Choix du modèle de base", level=2)
    para(doc,
        "Le modèle Phi-3-mini-4k-instruct de Microsoft a été retenu pour ses "
        "atouts : 3,8 milliards de paramètres seulement (donc déployable sur "
        "une carte graphique grand public), capacité multilingue native "
        "(français inclus), instruction-tuning de base déjà solide, contexte "
        "de 4096 tokens largement suffisant pour notre cas d'usage."
    )

    heading(doc, "5.2 Méthode LoRA", level=2)
    para(doc,
        "LoRA (Low-Rank Adaptation) consiste à figer les poids du modèle pré-"
        "entraîné et à n'entraîner que de petites matrices de bas rang ajoutées "
        "à chaque couche d'attention. Avec un rang de 16, nous n'adaptons "
        "qu'environ 1 % des paramètres totaux, ce qui réduit drastiquement "
        "le besoin en VRAM et le temps d'entraînement, tout en obtenant des "
        "performances proches d'un fine-tuning complet."
    )

    heading(doc, "5.3 Quantization 4-bit", level=2)
    para(doc,
        "Pour réduire encore le besoin en VRAM, nous chargeons le modèle de "
        "base en quantization 4-bit (bibliothèque bitsandbytes). Cette "
        "compression permet de faire tenir Phi-3-mini sur ~3 Go de VRAM, "
        "ce qui le rend exécutable sur des cartes graphiques modestes "
        "(GTX 1660, RTX 3050, etc.)."
    )

    heading(doc, "5.4 Hyperparamètres", level=2)
    table(doc,
        headers=["Paramètre", "Valeur", "Justification"],
        rows=[
            ["Modèle de base", "Phi-3-mini-4k-instruct", "Compromis taille / qualité"],
            ["Rang LoRA (r)", "16", "Standard pour des datasets petits"],
            ["Lora alpha", "32", "= 2 × r, valeur recommandée"],
            ["Lora dropout", "0,05", "Régularisation faible (peu de données)"],
            ["Modules cibles", "q_proj, k_proj, v_proj, o_proj",
             "Couches d'attention de Phi-3"],
            ["Learning rate", "2e-4", "Standard pour LoRA"],
            ["Batch size", "2", "Limité par la VRAM"],
            ["Gradient accumulation", "4", "Batch effectif = 8"],
            ["Époques", "3", "Suffit pour 65 exemples"],
            ["Optimiseur", "paged_adamw_8bit", "Optimisé mémoire"],
        ],
        widths_cm=[4.0, 4.5, 7.5],
    )

    heading(doc, "5.5 Architecture en microservice", level=2)
    para(doc,
        "Plutôt que d'embarquer le modèle dans le processus Django (ce qui "
        "alourdirait le démarrage et bloquerait la requête HTTP pendant "
        "l'inférence), nous avons fait le choix d'un microservice FastAPI "
        "dédié, exposé sur le port 8001. Cette architecture présente plusieurs "
        "avantages :"
    )
    bullets(doc, [
        "Découplage : le modèle peut être redéployé sans toucher au backend Django.",
        "Résilience : si le service IA tombe, l'application continue de fonctionner "
        "grâce au mécanisme de fallback du ChatService.",
        "Hébergement séparé : on peut placer le service sur une machine GPU "
        "dédiée si besoin.",
        "Mode mock : un mode de démonstration sans modèle (USE_MOCK=1) facilite "
        "le développement et les tests CI.",
    ])

    heading(doc, "5.6 Encapsulation : le ChatService", level=2)
    para(doc,
        "Côté Django, toute la complexité de l'appel HTTP, du formatage des "
        "messages et du fallback est encapsulée dans une classe ChatService. "
        "Les vues n'utilisent que sa méthode publique repondre(question, historique)."
    )
    code_block(doc,
        "class ChatService:\n"
        "    def __init__(self, ia_url=None, timeout=30):\n"
        "        self.ia_url = (ia_url or settings.IA_SERVICE_URL).rstrip('/')\n"
        "        self.timeout = timeout\n\n"
        "    def repondre(self, question, historique=None) -> str:\n"
        "        try:\n"
        "            return self._appeler_chat(question, historique or [])\n"
        "        except IAServiceError as exc:\n"
        "            logger.warning('IA injoignable : %s', exc)\n"
        "            return self._fallback(question)"
    )

    heading(doc, "5.7 Tests unitaires", level=2)
    para(doc,
        "Le ChatService est couvert par 8 tests unitaires utilisant unittest.mock "
        "pour simuler les réponses HTTP : succès simple, gestion de l'historique, "
        "timeout, erreur réseau, réponse vide, health online/offline. Cette "
        "couverture garantit que la logique de fallback fonctionne dans tous "
        "les cas d'erreur attendus."
    )

    heading(doc, "5.8 Concepts POO mobilisés (synthèse)", level=2)
    table(doc,
        headers=["Concept", "Illustration dans le projet"],
        rows=[
            ["Héritage", "Utilisateur(AbstractUser), templates extends base.html, "
             "InscriptionForm(UserCreationForm)"],
            ["Encapsulation", "Méthodes métier (peut_emprunter, est_disponible, "
             "valider, refuser), ChatService masquant l'HTTP"],
            ["Polymorphisme", "TextChoices (Role, EtatMateriel, StatutDemande), "
             "navbar adaptée au rôle"],
            ["Abstraction", "Serializers DRF, ViewSets génériques, ModelHolder "
             "isolant le modèle IA"],
            ["Composition", "Demande contient des LigneDemande, ConversationChat "
             "contient une liste de messages"],
        ],
        widths_cm=[3.5, 12.5],
    )

    # =====================================================================
    # CONCLUSION
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Conclusion générale", level=1)

    heading(doc, "Bilan du projet", level=2)
    para(doc,
        "Au terme de ce projet, nous avons conçu, développé et documenté une "
        "application web complète couvrant tout le cycle d'emprunt du matériel "
        "topographique de l'UFR SI. L'application met en œuvre un backend "
        "Django robuste, une API REST conforme aux standards, un frontend "
        "responsive avec carte interactive, et un assistant IA fine-tuné "
        "déployé en microservice."
    )
    para(doc,
        "Les concepts de programmation orientée objet enseignés en cours "
        "(héritage, encapsulation, polymorphisme, abstraction) ont été "
        "concrètement mobilisés et leur intérêt vérifié : le code reste "
        "lisible, testable, et extensible."
    )

    heading(doc, "Acquis personnels", level=2)
    para(doc,
        "Ce projet nous a permis de découvrir et d'approfondir de nombreux "
        "outils et méthodes : la modélisation UML, la méthode Merise, le "
        "framework Django et son ORM, la conception d'API REST, la "
        "cartographie web avec Leaflet, le fine-tuning de modèles de langage "
        "avec LoRA, les microservices FastAPI, et la collaboration en binôme "
        "via Git avec des branches de fonctionnalité."
    )
    para(doc,
        "Nous retenons aussi l'importance d'une conception soignée en amont : "
        "les heures investies dans les diagrammes UML et la modélisation des "
        "données ont été largement rentabilisées par un développement plus "
        "rapide et moins de refactoring."
    )

    heading(doc, "Perspectives", level=2)
    para(doc,
        "Plusieurs extensions sont envisageables pour faire évoluer la plateforme :"
    )
    bullets(doc, [
        "Notifications par e-mail lors des changements de statut d'une demande.",
        "Tableau de bord administrateur enrichi avec statistiques de l'usage "
        "du matériel, taux de retard, matériel le plus emprunté.",
        "Export CSV / Excel des demandes pour les rapports périodiques.",
        "Application mobile (PWA ou native) pour les demandes sur le terrain.",
        "Intégration PostGIS pour des requêtes spatiales avancées (matériel "
        "le plus utilisé dans telle zone).",
        "Évaluation rigoureuse du chatbot par un panel d'utilisateurs, avec "
        "métriques BLEU/ROUGE et taux de satisfaction.",
        "Hybridation RAG + fine-tuning : utiliser une recherche vectorielle "
        "pour les nouvelles fiches matériel, et le fine-tuning pour le ton "
        "et les procédures stables.",
    ])

    heading(doc, "Mot de la fin", level=2)
    para(doc,
        "Ce projet nous a confirmé l'intérêt de la programmation orientée "
        "objet pour structurer un système réel, et nous a donné le goût "
        "des projets logiciels complets, de la conception à la mise en "
        "service. Nous remercions encore une fois nos enseignants pour "
        "les bases solides qu'ils nous ont transmises."
    )

    # =====================================================================
    # RÉFÉRENCES
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Références", level=1)
    refs = [
        "Django Software Foundation. Django Documentation. https://docs.djangoproject.com",
        "Encode OSS Ltd. Django REST Framework. https://www.django-rest-framework.org",
        "Bootstrap. Bootstrap 5 Documentation. https://getbootstrap.com/docs/5.3/",
        "Volodymyr Agafonkin. Leaflet — an open-source JavaScript library for "
        "mobile-friendly interactive maps. https://leafletjs.com",
        "Hu, E. J., Shen, Y., Wallis, P., et al. (2021). LoRA: Low-Rank "
        "Adaptation of Large Language Models. arXiv:2106.09685.",
        "Microsoft Research. Phi-3 Technical Report. https://huggingface.co/"
        "microsoft/Phi-3-mini-4k-instruct",
        "Hugging Face. PEFT — Parameter-Efficient Fine-Tuning. "
        "https://huggingface.co/docs/peft",
        "OpenStreetMap contributors. https://www.openstreetmap.org",
        "Cours de Programmation Orientée Objet, UFR Sciences de l'Ingénieur, "
        "Université Iba Der Thiam de Thiès, 2025-2026.",
        "Cours de Topographie et Géodésie, L1-L2 Géomatique, UFR SI, 2024-2025.",
    ]
    for r in refs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(r)
        run.font.name = "Arial"; run.font.size = Pt(10)

    # =====================================================================
    # ANNEXES
    # =====================================================================
    doc.add_page_break()
    heading(doc, "Annexes", level=1)

    heading(doc, "Annexe A — Structure du dépôt", level=2)
    code_block(doc,
        "Projet POO/\n"
        "├── README.md\n"
        "├── backend/                 # Projet Django\n"
        "│   ├── config/              # settings, urls, wsgi\n"
        "│   ├── users/               # App utilisateurs\n"
        "│   ├── materiel/            # App matériels\n"
        "│   ├── emprunts/            # App demandes / restitutions\n"
        "│   ├── chatbot/             # App conversations IA\n"
        "│   ├── api/                 # Routeur REST central\n"
        "│   ├── pages/               # Vues HTML\n"
        "│   ├── fixtures/            # Données initiales JSON\n"
        "│   └── requirements.txt\n"
        "├── frontend/                # Templates et statics\n"
        "│   ├── templates/\n"
        "│   └── static/\n"
        "├── ia/                      # Module intelligence artificielle\n"
        "│   ├── corpus/              # Sources de connaissances\n"
        "│   ├── dataset/             # Paires Q/R + script de génération\n"
        "│   ├── training/            # Fine-tuning LoRA\n"
        "│   └── service/             # Microservice FastAPI\n"
        "└── docs/                    # Documentation\n"
        "    ├── conception/          # UML, MCD/MLD, architecture\n"
        "    ├── manuels/             # Manuel utilisateur\n"
        "    └── soutenance/          # Rapport et présentation"
    )

    heading(doc, "Annexe B — Stack technique synthétisée", level=2)
    table(doc,
        headers=["Couche", "Technologie", "Version"],
        rows=[
            ["Langage", "Python", "3.11+"],
            ["Backend", "Django", "5.0.6"],
            ["API REST", "Django REST Framework", "3.15.1"],
            ["Auth", "djangorestframework-simplejwt", "5.3.1"],
            ["Base de données", "SQLite", "intégré"],
            ["Frontend CSS", "Bootstrap", "5.3.3"],
            ["Frontend JS", "Vanilla", "—"],
            ["Cartographie", "Leaflet.js + OpenStreetMap", "1.9.4"],
            ["IA — modèle base", "Phi-3-mini-4k-instruct", "—"],
            ["IA — adaptation", "PEFT + LoRA", "0.10+"],
            ["IA — service", "FastAPI", "0.110+"],
            ["Versionnement", "Git + GitHub", "—"],
        ],
        widths_cm=[4.0, 8.0, 4.0],
    )

    heading(doc, "Annexe C — Lien vers le code source", level=2)
    para(doc,
        "Le code source complet, l'historique Git, les issues et la documentation "
        "technique détaillée sont disponibles sur le dépôt GitHub du projet :"
    )
    p = doc.add_paragraph()
    r = p.add_run("https://github.com/<pseudo>/gestion-emprunt-materiel-ufr")
    r.font.name = "Consolas"; r.font.size = Pt(11); r.font.color.rgb = BLUE

    doc.add_paragraph()
    styled(doc, "— Fin du rapport —",
           align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10, color=GREY)

    # ---------- SAUVEGARDE ----------
    doc.save(str(OUTPUT))
    print(f"[OK] Rapport généré : {OUTPUT}")


if __name__ == "__main__":
    build()
