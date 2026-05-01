"""
build_manuel.py — Génère le manuel utilisateur en .docx à partir du markdown.

Préserve les accents français nativement (UTF-8 partout).
Dépendance : python-docx (pip install python-docx)

Usage :
    python build_manuel.py
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).parent
OUTPUT = ROOT / "manuel_utilisateur.docx"


# ---------------------------------------------------------------------------
# Helpers de mise en forme
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color_hex: str) -> None:
    """Applique une couleur de fond à une cellule (sans noircir)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str = "CCCCCC") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_page_number_field(paragraph) -> None:
    """Insère un champ « Page X » dans un footer/header."""
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)

    # Champ PAGE
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run2 = paragraph.add_run()
    run2.font.name = "Arial"
    run2.font.size = Pt(9)
    run2._r.append(fld_char1)
    run2._r.append(instr_text)
    run2._r.append(fld_char2)


def add_table_of_contents(doc) -> None:
    """Insère un champ TOC (à rafraîchir manuellement dans Word : F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:t")
    fld_char3.text = "Mettez à jour la table des matières (clic droit > Mettre à jour)."
    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    run._r.append(fld_char4)


def add_styled_paragraph(doc, text: str, style: str | None = None,
                          align=None, bold=False, color=None, size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Arial"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]],
              col_widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False

    # En-tête
    for j, txt in enumerate(headers):
        cell = table.cell(0, j)
        cell.width = Cm(col_widths_cm[j])
        set_cell_shading(cell, "D5E8F0")
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = True

    # Lignes
    for i, row in enumerate(rows, start=1):
        for j, txt in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Cm(col_widths_cm[j])
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(txt)
            run.font.name = "Arial"
            run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Construction du document
# ---------------------------------------------------------------------------
def build():
    doc = Document()

    # Styles globaux
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # Marges
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = True

    # ---------- COUVERTURE ----------
    for _ in range(6):
        doc.add_paragraph()
    add_styled_paragraph(
        doc, "Manuel utilisateur",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=32,
        color=RGBColor(0x0D, 0x6E, 0xFD),
    )
    add_styled_paragraph(
        doc, "Plateforme de gestion d'emprunt du matériel",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20,
    )
    add_styled_paragraph(
        doc, "UFR Sciences de l'Ingénieur",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=16,
    )
    add_styled_paragraph(
        doc, "Université Iba Der Thiam de Thiès",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
    )
    for _ in range(6):
        doc.add_paragraph()
    add_styled_paragraph(
        doc, "Projet de Programmation Orientée Objet",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12,
    )
    add_styled_paragraph(
        doc, "Licence 2 Géomatique",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
    )
    add_styled_paragraph(
        doc, "Bineta Elimane Hanne  &  Aminata Kounta",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12,
    )
    for _ in range(2):
        doc.add_paragraph()
    add_styled_paragraph(
        doc, "Version 1.0 — avril 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
        color=RGBColor(0x6C, 0x75, 0x7D),
    )

    # ---------- TOC ----------
    doc.add_page_break()
    h = doc.add_heading("Table des matières", level=1)
    h.runs[0].font.name = "Arial"
    add_table_of_contents(doc)

    # ---------- HEADER / FOOTER pour les pages de contenu ----------
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run("Manuel utilisateur — UFR SI")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(fp)

    # Première page (couverture) : pas de header/footer
    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""

    # ---------- CONTENU ----------
    # Section 1
    doc.add_page_break()
    doc.add_heading("1. Bienvenue", level=1)
    doc.add_paragraph(
        "Cette plateforme permet aux étudiants, enseignants et techniciens de "
        "l'UFR Sciences de l'Ingénieur de gérer en ligne l'emprunt du matériel "
        "topographique et géodésique : stations totales, récepteurs GNSS, "
        "niveaux, accessoires, équipements informatiques mobiles."
    )
    doc.add_paragraph("Vous y trouverez :")
    for item in [
        "Un catalogue interactif du matériel disponible",
        "Un formulaire de demande avec géolocalisation sur carte",
        "Un suivi en temps réel du statut de vos demandes",
        "Un assistant IA pour répondre à vos questions techniques",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "L'objectif est de remplacer la gestion papier par un système traçable, "
        "transparent et accessible 24 h / 24."
    )

    # Section 2
    doc.add_page_break()
    doc.add_heading("2. Premier accès", level=1)
    doc.add_heading("2.1 Créer un compte", level=2)
    for item in [
        "Ouvrez votre navigateur et allez sur l'adresse de la plateforme "
        "(par défaut : http://127.0.0.1:8000).",
        "Cliquez sur le bouton « S'inscrire » en haut à droite.",
        "Remplissez le formulaire avec vos prénom, nom, e-mail UFR, filière, niveau, "
        "téléphone (optionnel), nom d'utilisateur et un mot de passe robuste "
        "(au moins 8 caractères, avec majuscules, chiffres et caractères spéciaux).",
        "Cliquez sur « Créer mon compte ».",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph(
        "Votre compte est créé immédiatement avec le rôle Étudiant. "
        "Vous êtes connecté automatiquement."
    )
    doc.add_heading("2.2 Se connecter", level=2)
    for item in [
        "Cliquez sur « Connexion » en haut à droite.",
        "Entrez votre nom d'utilisateur et votre mot de passe.",
        "Cliquez sur « Se connecter ».",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_heading("2.3 Se déconnecter", level=2)
    doc.add_paragraph(
        "Cliquez sur votre prénom en haut à droite, puis sur « Se déconnecter »."
    )

    # Section 3
    doc.add_page_break()
    doc.add_heading("3. Catalogue du matériel", level=1)
    doc.add_heading("3.1 Consulter le catalogue", level=2)
    doc.add_paragraph(
        "Cliquez sur « Catalogue » dans la barre de navigation. Vous voyez une grille "
        "de toutes les références disponibles à l'UFR. Pour chaque matériel, vous trouvez :"
    )
    for item in [
        "Nom et modèle",
        "Catégorie (Station totale, GNSS, Niveau, Accessoire, etc.)",
        "État (Disponible, Emprunté, Maintenance, Hors service)",
        "Stock (quantité disponible / quantité totale)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("3.2 Filtrer par catégorie", level=2)
    doc.add_paragraph(
        "Au-dessus de la grille, des boutons de filtre permettent d'afficher seulement "
        "les matériels d'une catégorie. Cliquez sur « Toutes » pour revenir à la vue complète."
    )
    doc.add_heading("3.3 Comprendre les états", level=2)
    add_table(
        doc,
        headers=["Pastille", "Signification"],
        rows=[
            ["Vert « Disponible »", "Au moins une unité libre, vous pouvez l'emprunter"],
            ["Jaune « Emprunté »", "Toutes les unités sont actuellement sorties"],
            ["Bleu « Maintenance »", "En réparation, indisponible temporairement"],
            ["Gris « Hors service »", "Définitivement indisponible"],
        ],
        col_widths_cm=[5.0, 11.0],
    )

    # Section 4
    doc.add_page_break()
    doc.add_heading("4. Faire une demande d'emprunt", level=1)
    doc.add_heading("4.1 Accéder au formulaire", level=2)
    doc.add_paragraph(
        "Cliquez sur « Nouvelle demande » dans la barre de navigation, ou sur le gros "
        "bouton vert depuis la page « Catalogue »."
    )
    doc.add_heading("4.2 Remplir les informations de période", level=2)
    doc.add_paragraph("Dans la première carte « Période d'emprunt » :")
    for item in [
        "Choisissez la date de début : jour où vous récupérerez le matériel.",
        "Choisissez la date de fin : jour où vous le rendrez. La date de fin est "
        "automatiquement contrainte à être supérieure ou égale à la date de début.",
        "Saisissez le motif de la sortie. Soyez précis : un motif vague risque de "
        "faire refuser votre demande.",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_heading("4.3 Localiser votre zone de travail", level=2)
    doc.add_paragraph("Dans la carte « Localisation » :")
    for item in [
        "La carte est centrée sur Thiès par défaut.",
        "Cliquez à l'endroit exact où vous comptez utiliser le matériel. "
        "Un marqueur rouge apparaît.",
        "Vous pouvez glisser-déposer ce marqueur pour ajuster sa position.",
        "Les champs Latitude et Longitude se remplissent automatiquement.",
        "Saisissez optionnellement un nom de lieu (ex : « Campus UFR SI », "
        "« Carrière de Pout »).",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph(
        "Cette géolocalisation aide les enseignants à comprendre votre besoin "
        "et permet à l'UFR de cartographier l'usage du matériel."
    )
    doc.add_heading("4.4 Sélectionner le matériel", level=2)
    for item in [
        "Faites défiler la liste des équipements disponibles.",
        "Cochez la case devant chaque matériel souhaité.",
        "Le champ quantité s'active automatiquement. Saisissez la quantité voulue "
        "(limitée au stock disponible affiché).",
        "Le compteur en haut à droite indique combien d'articles vous avez sélectionnés.",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_heading("4.5 Soumettre", level=2)
    doc.add_paragraph(
        "Cliquez sur le bouton vert « Soumettre ma demande ». Vous êtes redirigé "
        "vers « Mes demandes » où votre nouvelle demande apparaît avec le statut "
        "« En attente »."
    )
    doc.add_heading("4.6 Délais et bonnes pratiques", level=2)
    for item in [
        "Soumettez votre demande au moins 48 heures à l'avance.",
        "Pour un TP encadré ou un projet important : une semaine d'avance.",
        "En période d'examens, les ressources sont très sollicitées : prévoyez large.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Section 5
    doc.add_page_break()
    doc.add_heading("5. Suivre mes demandes", level=1)
    doc.add_heading("5.1 Vue d'ensemble", level=2)
    doc.add_paragraph(
        "Cliquez sur « Mes emprunts » dans la barre de navigation. Vous voyez "
        "un tableau récapitulant toutes vos demandes, de la plus récente à la plus ancienne."
    )
    doc.add_heading("5.2 Comprendre les statuts", level=2)
    add_table(
        doc,
        headers=["Statut", "Signification", "Action attendue"],
        rows=[
            ["En attente", "Soumise, attente de validation", "Patientez (48 h ouvrées max)"],
            ["Approuvée", "Validée, prête à retirer", "Rendez-vous au laboratoire"],
            ["Refusée", "Refusée par le validateur", "Lisez le motif, refaites une demande"],
            ["En cours", "Matériel retiré, sortie en cours", "Rendez le matériel à la date prévue"],
            ["Restituée", "Matériel rendu, dossier clos", "Rien à faire"],
            ["Annulée", "Annulée par vous-même", "Rien à faire"],
        ],
        col_widths_cm=[3.5, 6.5, 6.0],
    )
    doc.add_heading("5.3 Voir les détails d'une demande", level=2)
    doc.add_paragraph(
        "Cliquez sur le bouton « Détails » au bout de chaque ligne. Une zone "
        "se déplie en dessous avec le motif complet, la liste des articles avec "
        "quantités, et le lieu avec les coordonnées GPS."
    )
    doc.add_heading("5.4 Annuler une demande", level=2)
    doc.add_paragraph(
        "Tant qu'une demande est en statut « En attente », vous pouvez l'annuler. "
        "Une fois approuvée, vous devez contacter directement l'enseignant ou le technicien."
    )

    # Section 6
    doc.add_page_break()
    doc.add_heading("6. Restitution", level=1)
    doc.add_heading("6.1 Avant de venir rendre le matériel", level=2)
    doc.add_paragraph("Vérifiez que :")
    for item in [
        "Les batteries sont rechargées (utilisez les chargeurs fournis).",
        "Le matériel est propre : essuyez la poussière, vérifiez les optiques.",
        "Les trépieds sont pliés, sangles serrées.",
        "Les prismes sont rangés dans leur étui.",
        "Tous les accessoires empruntés sont présents (cordons, télécommandes, mires, etc.).",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("6.2 À la restitution", level=2)
    doc.add_paragraph(
        "Le technicien vérifie chaque article. Trois cas possibles pour chaque pièce :"
    )
    for item in [
        "Bon état : le matériel revient en stock, rien à signaler.",
        "Endommagé : une fiche de maintenance est ouverte. Selon la nature et la cause "
        "du dommage, vous pouvez être amené à participer aux frais.",
        "Perdu : vous devez rembourser le matériel au prix de remplacement, après "
        "vérification (déclaration de perte, recherche).",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("6.3 En cas de retard", level=2)
    doc.add_paragraph(
        "Un rappel automatique vous est envoyé. Sans justificatif valable, vous pouvez "
        "être suspendu temporairement de la plateforme et perdre la priorité sur vos "
        "prochaines demandes. Prévenez toujours le technicien dès que vous savez que "
        "vous serez en retard."
    )

    # Section 7
    doc.add_page_break()
    doc.add_heading("7. Assistant IA", level=1)
    doc.add_heading("7.1 Accéder à l'assistant", level=2)
    doc.add_paragraph(
        "Cliquez sur « Assistant IA » dans la barre de navigation. Une fenêtre "
        "de discussion s'ouvre."
    )
    doc.add_heading("7.2 Poser une question", level=2)
    for item in [
        "Tapez votre question dans le champ en bas (par exemple : « Quelle est la "
        "précision de la station Leica TS06 ? »).",
        "Cliquez sur « Envoyer » ou pressez Entrée.",
        "L'assistant répond en quelques secondes.",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_heading("7.3 Que peut-il faire ?", level=2)
    doc.add_paragraph("L'assistant est entraîné sur :")
    for item in [
        "Les fiches techniques du matériel UFR (Leica TS06, Topcon CTS-112 R4, "
        "CHC i50/i73, Garmin, niveaux, etc.).",
        "Les procédures d'emprunt et de restitution.",
        "Les notions de topographie et géodésie (angles, distances, nivellement, GNSS, "
        "projections UTM Sénégal).",
        "Les questions fréquentes des étudiants.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("7.4 Que ne peut-il pas faire ?", level=2)
    for item in [
        "Faire vos calculs de TP à votre place.",
        "Répondre à des questions hors-périmètre (autres cours, vie universitaire en général).",
        "Remplacer un enseignant pour des questions de fond complexes.",
        "Modifier vos demandes ou réserver du matériel pour vous.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("7.5 Conversations passées", level=2)
    doc.add_paragraph(
        "La barre latérale gauche liste vos conversations passées. Cliquez sur "
        "l'une d'elles pour reprendre la discussion. Cliquez sur « Nouvelle » "
        "pour démarrer une discussion vierge."
    )

    # Section 8
    doc.add_page_break()
    doc.add_heading("8. Profils et rôles", level=1)
    doc.add_paragraph(
        "La plateforme distingue quatre rôles, avec des permissions différentes."
    )
    add_table(
        doc,
        headers=["Rôle", "Comment l'obtenir", "Permissions principales"],
        rows=[
            ["Étudiant", "Inscription sur la plateforme",
             "Catalogue, demandes, suivi, chatbot"],
            ["Enseignant", "Attribution par administrateur",
             "Idem étudiant + valider/refuser les demandes"],
            ["Technicien", "Attribution par administrateur",
             "Idem étudiant + sorties, restitutions, maintenance"],
            ["Administrateur", "Attribution par administrateur",
             "Tous droits + gestion utilisateurs et catalogue (/admin/)"],
        ],
        col_widths_cm=[3.0, 5.5, 7.5],
    )

    # Section 9
    doc.add_page_break()
    doc.add_heading("9. Foire aux questions", level=1)

    faq = [
        ("J'ai oublié mon mot de passe.",
         "Contactez le technicien ou l'administrateur de l'UFR. Ils peuvent "
         "réinitialiser votre mot de passe manuellement. Une fonction « mot de "
         "passe oublié » par e-mail sera ajoutée prochainement."),
        ("Combien de matériels puis-je emprunter en même temps ?",
         "Pas de limite stricte, mais votre demande doit être justifiée. Le "
         "validateur peut refuser ou ajuster une demande qui mobilise trop de matériel."),
        ("Puis-je faire une demande pour mon binôme et moi ?",
         "Oui. Une seule demande, faite par le chef de binôme, avec les deux "
         "noms dans le motif."),
        ("Que se passe-t-il si du matériel tombe en panne pendant ma sortie ?",
         "Si vous prouvez une utilisation conforme, vous n'êtes pas responsable. "
         "Signalez la panne au technicien à la restitution. Une fiche de "
         "maintenance est ouverte automatiquement."),
        ("J'ai perdu un prisme, dois-je rembourser ?",
         "Oui, sauf cas de force majeure prouvé (vol avec dépôt de plainte). "
         "Le prix de remplacement vous sera communiqué par le technicien."),
        ("Le chatbot peut-il faire mes calculs de polygonale ?",
         "Il peut expliquer la méthode et les formules. Pour le calcul lui-même, "
         "utilisez Excel ou un logiciel dédié comme Covadis."),
        ("Mes données sont-elles confidentielles ?",
         "Oui. Seuls vous, l'enseignant validateur et le technicien voient vos "
         "demandes. Aucune donnée n'est partagée avec des tiers."),
    ]
    for q, r in faq:
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"Q : {q}")
        run_q.bold = True
        run_q.font.name = "Arial"
        p_r = doc.add_paragraph()
        run_r = p_r.add_run(f"R : {r}")
        run_r.font.name = "Arial"
        doc.add_paragraph()

    # Section 10
    doc.add_page_break()
    doc.add_heading("10. Contact et support", level=1)
    doc.add_paragraph(
        "Pour signaler un bug, suggérer une amélioration, ou demander de l'aide :"
    )
    for item in [
        "Auteures du projet : Bineta Elimane Hanne, Aminata Kounta",
        "Dépôt GitHub : ouvrir une issue sur le dépôt du projet",
        "Sur place : laboratoire de topographie de l'UFR SI",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()
    add_styled_paragraph(
        doc, "Bonne utilisation !",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14,
        color=RGBColor(0x0D, 0x6E, 0xFD),
    )
    doc.add_paragraph()
    add_styled_paragraph(
        doc, "Manuel utilisateur — version 1.0 — avril 2026.",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=9,
        color=RGBColor(0x6C, 0x75, 0x7D),
    )

    # Sauvegarde
    doc.save(str(OUTPUT))
    print(f"[OK] Document généré : {OUTPUT}")


if __name__ == "__main__":
    build()
